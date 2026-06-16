from __future__ import annotations

import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from backend.app.core.errors import UnsupportedDocumentError


@dataclass
class ParsedDocument:
    text: str
    page_texts: list[tuple[int, str]]
    warnings: list[str]


def parse_txt(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return ParsedDocument(text=text, page_texts=[(1, text)], warnings=[])


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise UnsupportedDocumentError("python-docx is required for DOCX parsing.") from exc
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return ParsedDocument(text=text, page_texts=[(1, text)], warnings=[])


def _parse_pdf_pymupdf(path: Path) -> ParsedDocument:
    import fitz  # type: ignore

    page_texts: list[tuple[int, str]] = []
    warnings: list[str] = []
    with fitz.open(path) as pdf:
        for index, page in enumerate(pdf, start=1):
            page_text = page.get_text("text") or ""
            if not page_text.strip():
                warnings.append(f"Page {index} had no extractable text.")
            page_texts.append((index, page_text))
    return ParsedDocument(
        text="\n".join(page_text for _, page_text in page_texts),
        page_texts=page_texts,
        warnings=warnings,
    )


def _parse_pdf_pdfplumber(path: Path) -> ParsedDocument:
    import pdfplumber  # type: ignore

    page_texts: list[tuple[int, str]] = []
    warnings: list[str] = ["PyMuPDF failed; used pdfplumber fallback."]
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                warnings.append(f"Page {index} had no extractable text.")
            page_texts.append((index, page_text))
    return ParsedDocument(
        text="\n".join(page_text for _, page_text in page_texts),
        page_texts=page_texts,
        warnings=warnings,
    )


def _text_quality(text: str) -> float:
    if not text:
        return 0.0
    useful = sum(1 for char in text if char.isprintable() or char in "\n\t")
    letters = sum(1 for char in text if char.isalpha())
    return min(useful / max(len(text), 1), letters / max(len(text), 1) * 4)


def _parse_pdf_pdftotext(path: Path) -> ParsedDocument:
    if not shutil.which("pdftotext"):
        raise UnsupportedDocumentError("pdftotext is unavailable.")
    result = subprocess.run(
        ["pdftotext", "-layout", str(path), "-"],
        check=True,
        capture_output=True,
        text=True,
        timeout=30,
    )
    pages = result.stdout.split("\f")
    page_texts = [
        (index, page_text)
        for index, page_text in enumerate(pages, start=1)
        if page_text.strip()
    ]
    return ParsedDocument(
        text="\n".join(page_text for _, page_text in page_texts),
        page_texts=page_texts or [(1, result.stdout)],
        warnings=["Used pdftotext fallback because embedded PDF text looked garbled."],
    )


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        parsed = _parse_pdf_pymupdf(path)
        if _text_quality(parsed.text) >= 0.45:
            return parsed
        try:
            fallback = _parse_pdf_pdftotext(path)
            if _text_quality(fallback.text) > _text_quality(parsed.text):
                return fallback
        except Exception:
            return parsed
        return parsed
    except Exception:
        try:
            parsed = _parse_pdf_pdfplumber(path)
            if _text_quality(parsed.text) >= 0.45:
                return parsed
            try:
                fallback = _parse_pdf_pdftotext(path)
                if _text_quality(fallback.text) > _text_quality(parsed.text):
                    return fallback
            except Exception:
                return parsed
            return parsed
        except Exception as exc:
            try:
                return _parse_pdf_pdftotext(path)
            except Exception:
                raise UnsupportedDocumentError(
                    "PDF parsing failed with PyMuPDF, pdfplumber, and pdftotext. OCR may be required."
                ) from exc


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return parse_txt(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    raise UnsupportedDocumentError(f"Unsupported file type: {suffix or 'unknown'}")
